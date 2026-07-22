from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Axnovus_Care_Agent_DPIIT_Startup_India_Investor_Document.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
CYAN = RGBColor(12, 151, 179)
MUTED = RGBColor(90, 103, 120)
GRAY_FILL = "F4F6F9"
BLUE_FILL = "E8EEF5"
LIGHT_CYAN = "EAF8FB"
BORDER = "D8DEE8"
BLACK = RGBColor(0, 0, 0)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_paragraph_border_bottom(paragraph, color="00A7C8", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_run(run, size=None, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, size=None, color=None, bold=None, italic=None, align=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run(run, size=16 if level == 1 else 13 if level == 2 else 12, color=BLUE if level < 3 else NAVY, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, size=10.6, color=BLACK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, size=10.6, color=BLACK)
    return p


def add_callout(doc, label, text, fill=LIGHT_CYAN):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_table_borders(table, color="B7DDE8", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_run(r, size=10.5, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run(r2, size=10.3, color=BLACK)
    add_para(doc, "", after=4)
    return table


def add_kv_table(doc, rows, widths=(2300, 7060), header=None):
    table = doc.add_table(rows=1 if header else 0, cols=2)
    if header:
        hdr = table.rows[0]
        hdr.cells[0].text = header[0]
        hdr.cells[1].text = header[1]
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    set_table_width(table, list(widths))
    set_table_borders(table)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if header and i == 0:
                set_cell_shading(cell, BLUE_FILL)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run(run, size=9.7, color=NAVY if j == 0 else BLACK, bold=(j == 0 or (header and i == 0)))
    add_para(doc, "", after=4)
    return table


def add_matrix(doc, headers, rows, widths, header_fill=BLUE_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    set_table_width(table, widths)
    set_table_borders(table)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                set_cell_shading(cell, header_fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run(run, size=9.1 if len(headers) >= 4 else 9.5, color=NAVY if i == 0 else BLACK, bold=(i == 0))
    add_para(doc, "", after=4)
    return table


def add_page(doc, title, subtitle=None):
    doc.add_page_break()
    add_heading(doc, title, 1)
    if subtitle:
        add_para(doc, subtitle, size=10.8, color=MUTED, italic=True, after=10)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [("Title", 24, NAVY), ("Subtitle", 12, MUTED)]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.text = "AXNOVUS | Healthcare AI Agent DPIIT Recognition Support Document"
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        set_run(run, size=8.8, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.text = "Prepared for Startup India DPIIT recognition and investor diligence | Founder review draft"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run(run, size=8.5, color=MUTED)


def build_doc():
    doc = Document()
    configure_document(doc)

    # Cover page
    add_para(doc, "AXNOVUS", size=13, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_para(doc, "AI Agents and Intelligent Products", size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=70)
    add_para(doc, "Axnovus Care Agent", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "DPIIT Startup India Recognition", size=16, color=BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Support Document and Investor Brief", size=14, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(
        doc,
        "A concise, submission-ready business and innovation narrative for an India-context healthcare AI triage, routing, appointment, doctor-workflow, and prescription-support platform.",
        size=11,
        color=BLACK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=30,
    )
    rule = add_para(doc, "", after=18)
    set_paragraph_border_bottom(rule, color="00A7C8", size="14")
    add_kv_table(
        doc,
        [
            ("Purpose", "Position Axnovus Care Agent as an innovation-led, scalable healthcare SaaS/AI startup with employment, wealth creation, and public-health value."),
            ("Use", "Attach with Startup India application materials and share as a first-pass investor, accelerator, grant, or hospital-partnership overview after legal and financial facts are updated."),
            ("Version", "Founder review draft | June 2026"),
            ("Confidentiality", "Private and confidential business document"),
        ],
        widths=(1900, 7460),
    )

    add_page(doc, "Executive Recognition Snapshot")
    add_para(
        doc,
        "Axnovus Care Agent is proposed as an AI-enabled healthcare coordination platform for Indian hospital OPD and care-navigation workflows. It accepts symptoms through typed text or Hindi speech, considers patient context and uploaded reports, generates possible conditions with targeted follow-up questions, routes patients to appropriate hospital specialists, supports appointment booking, and gives doctors a structured workspace for case review, additional report requests, prescriptions, and medicine option comparison based on salt composition.",
        size=10.8,
        after=8,
    )
    add_callout(
        doc,
        "DPIIT fit",
        "The venture is software-led, AI-enabled, modular, scalable across hospitals and clinics, and capable of creating skilled jobs in product, healthcare operations, clinical review, customer success, integrations, and hospital partnerships.",
    )
    add_heading(doc, "Recognition logic", 2)
    for item in [
        "Innovation: combines multilingual symptom intake, LLM-primary triage, emergency guardrails, report-aware reasoning, dynamic follow-up, specialist routing, doctor handoff, and feedback-driven local case learning.",
        "Scalability: deployable as a hospital website widget, clinic OPD assistant, enterprise hospital module, insurance/care-navigation layer, or standalone patient-doctor workflow.",
        "Employment potential: creates direct roles in engineering, clinical operations, implementation, support, data governance, model evaluation, and hospital success.",
        "Wealth creation: supports subscription SaaS, hospital licensing, integration fees, analytics, marketplace/referral workflows, and clinical operations services.",
        "Public value: improves early care navigation, reduces front-desk ambiguity, supports Indian languages, and creates structured digital handoffs for doctors.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Founder action before submission", 2)
    for item in [
        "Replace bracketed legal fields with entity name, incorporation date, CIN/LLPIN, registered office, PAN, GST, founder details, and shareholding where required.",
        "Attach incorporation certificate, PAN, authorization letter, product screenshots, demo link, architecture note, founder profiles, and any pilot/customer evidence.",
        "Confirm current DPIIT recognition criteria on the official Startup India portal before upload, including entity form, incorporation age, turnover, and non-reconstruction conditions.",
    ]:
        add_bullet(doc, item)

    add_page(doc, "Startup Identity and Eligibility")
    add_kv_table(
        doc,
        [
            ("Startup / Brand", "Axnovus Care Agent by AxNovus / [Legal entity name to be inserted]"),
            ("Registered entity", "[Private Limited Company / LLP / Registered Partnership / eligible entity category as applicable]"),
            ("Incorporation details", "[Date of incorporation], [CIN/LLPIN/Registration number], [Registered office address]"),
            ("Founders / directors", "[Founder names], [roles], [relevant healthcare/AI/business experience]"),
            ("Sector classification", "HealthTech, AI/SaaS, hospital workflow automation, digital health, care navigation, clinical operations intelligence"),
            ("Recognition objective", "DPIIT Startup India recognition for innovation, scalability, employment generation, and wealth creation potential"),
        ],
    )
    add_heading(doc, "Eligibility assertions to include", 2)
    for item in [
        "The applicant entity is incorporated in India and is within the applicable Startup India recognition period from incorporation.",
        "The entity has not exceeded the applicable annual turnover threshold under current Startup India eligibility rules.",
        "The entity has not been formed by splitting up or reconstructing an existing business.",
        "The startup is working toward innovation, development, deployment, or commercialization of a product/process/service with high scalability and wealth/employment generation potential.",
        "Healthcare claims, customer traction, pilots, certifications, and clinical review statements must be limited to facts for which evidence is attached.",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Important compliance note",
        "This support document is not legal advice. DPIIT criteria and healthcare/data obligations should be verified against the official Startup India portal and professional advisors before final filing.",
        fill=GRAY_FILL,
    )

    add_page(doc, "Problem and Market Need")
    add_para(
        doc,
        "India's healthcare access problem is not only doctor availability. A major operational gap exists before the doctor sees the patient: symptoms are described in mixed language, patients do not know the right department, front desks make quick routing decisions with limited context, reports are not always considered early, and doctors receive inconsistent case histories. This creates unnecessary waiting, repeated questioning, avoidable misrouting, and weak continuity between digital intake and physical consultation.",
        size=10.8,
    )
    add_heading(doc, "Pain points", 2)
    for item in [
        "Unstructured intake: symptoms arrive as English, Hindi, Hinglish, abbreviations, voice notes, report files, or incomplete patient descriptions.",
        "Static triage tools: fixed rules and generic questions often ask the same follow-up questions for very different cases.",
        "Poor routing: patients often choose the wrong OPD department, causing delays, extra cost, and lower hospital satisfaction.",
        "Doctor handoff gap: physicians need symptoms, duration, red flags, reports, and prior answers in one compact view, not scattered notes.",
        "No learning loop: doctor corrections, prescriptions, report requests, and patient outcomes are rarely captured in a structured dataset for future improvement.",
        "Trust and safety risk: healthcare AI must be conservative, auditable, consent-aware, and clinically reviewed before it influences real care.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Why now", 2)
    for item in [
        "Indian hospitals and clinics are digitizing OPD, appointments, teleconsults, and patient engagement workflows.",
        "LLMs and structured-output APIs make broader symptom interpretation possible, including Hindi and mixed-language intake, when wrapped with safety guardrails.",
        "Patients increasingly expect self-service booking, report upload, prescription access, and medicine-price transparency.",
        "Hospitals need workflow automation that augments doctors and staff rather than adding another disconnected chatbot.",
    ]:
        add_bullet(doc, item)

    add_page(doc, "Axnovus Care Agent Solution")
    add_para(
        doc,
        "Axnovus Care Agent is designed as a configurable healthcare AI workflow layer, not a single static triage page. The product starts with patient intake and triage, then extends into doctor review, appointment booking, treatment capture, prescription upload, and medicine option comparison. Each module is separate enough to be embedded into hospital websites, clinic applications, care-navigation products, or enterprise health platforms.",
        size=10.8,
    )
    add_matrix(
        doc,
        ["Module", "What it does", "Customer value"],
        [
            ("Patient intake", "Typed symptoms, Hindi speech input, patient details, report upload, consent-ready case creation.", "Faster self-service intake with local language support."),
            ("AI triage", "LLM-primary symptom analysis, possible conditions, emergency red flags, and targeted follow-up questions.", "Broader coverage than static rules while retaining safety limits."),
            ("Routing and booking", "Maps likely condition and urgency to specialist, hospital, doctor, and slot options.", "Reduces misrouting and converts triage into an appointment."),
            ("Doctor workspace", "Shows booked patients, symptoms, follow-up answers, reports, AI summary, and doctor input fields.", "Improves consultation readiness and continuity."),
            ("Treatment and prescription", "Stores doctor notes, additional report requests, prescriptions, and patient-facing treatment details.", "Creates one structured case record after the physical visit."),
            ("Medicine options", "Parses prescription text for salt composition and compares configured brand/generic price options.", "Supports patient transparency and pharmacy integration readiness."),
        ],
        [1800, 4300, 3260],
    )
    add_heading(doc, "Current prototype evidence", 2)
    for item in [
        "Static Docker-runnable prototype with patient and doctor login paths.",
        "Hindi speech input through browser speech recognition where supported.",
        "Configurable LLM model selection with local-rule fallback.",
        "Report upload parsing, appointment slot matching, prescription capture, and salt-based medicine matching modules.",
        "Browser-local feedback dataset preparation using thumbs-up/down review signals.",
    ]:
        add_bullet(doc, item)

    add_page(doc, "Innovation and Differentiation")
    add_para(
        doc,
        "The innovation is not that the product asks symptoms. The differentiator is the full care-navigation loop: multilingual intake, model-backed reasoning, clinician-safe boundaries, interactive follow-up, specialist routing, appointment conversion, doctor review, prescription capture, and feedback learning. This creates a data and workflow asset that becomes more valuable as hospitals use it.",
        size=10.8,
    )
    add_matrix(
        doc,
        ["Differentiator", "Why it matters", "Defensibility"],
        [
            ("LLM-primary triage", "Covers broader symptom combinations than deterministic rules, especially mixed-language patient descriptions.", "Model prompts, structured outputs, clinical review sets, and local case datasets improve over time."),
            ("Report-aware reasoning", "Uploaded reports can influence triage context instead of being ignored until consultation.", "OCR, lab extraction, abnormal-range logic, and report summaries become domain-specific assets."),
            ("Doctor feedback loop", "Doctors can add follow-up questions, report requests, prescriptions, and corrections.", "Clinician-reviewed case repository can support retrieval, evaluation, and offline fallback after sufficient volume."),
            ("Configurable modules", "Symptoms, doctors, hospitals, medicines, models, and routing can be swapped for different hospital deployments.", "Reusable platform components reduce implementation cost across customers."),
            ("Safety-by-design", "Emergency red flags, non-diagnosis positioning, audit trails, consent, and review workflows reduce irresponsible automation risk.", "Trust and compliance can become a sales advantage in healthcare."),
        ],
        [2200, 3700, 3460],
    )
    add_callout(
        doc,
        "Critical positioning",
        "The product should be presented as a triage and care-routing assistant, not an autonomous diagnostic system. DPIIT and investors will respond better to a credible, governed healthcare workflow than to exaggerated medical-AI claims.",
    )

    add_page(doc, "Technology Architecture")
    add_para(
        doc,
        "The architecture should be presented as a staged system so evaluators can see current feasibility and long-term product depth. The current prototype is intentionally modular: configuration files hold clinical/routing content, service files handle stores and domain functions, and UI orchestration is separate from reusable engines.",
        size=10.8,
    )
    add_kv_table(
        doc,
        [
            ("Input layer", "Symptom text, Hindi/Hinglish speech, demographics, uploaded reports, patient consent, and prior case context."),
            ("Reasoning layer", "LLM structured outputs, local safety rules, configurable disease/routing knowledge base, follow-up planner, and confidence handling."),
            ("Workflow layer", "Patient stages, appointment booking, doctor queue, doctor notes, report requests, prescription capture, medicine comparison, and feedback."),
            ("Data layer", "Secure backend database in production, audit logs, de-identified training/evaluation sets, versioned model prompts, and clinical review status."),
            ("Integration layer", "Hospital HIS/EMR, appointment APIs, ABDM/ABHA where appropriate, pharmacy pricing/inventory APIs, OCR/lab extraction, and notification systems."),
            ("Trust layer", "Role-based access, consent, retention/deletion policy, DPDP-aligned privacy controls, clinician review, red-flag escalation, and model monitoring."),
        ],
    )
    add_heading(doc, "Development roadmap", 2)
    for item in [
        "MVP: patient/doctor login, intake, triage, follow-up, routing, appointment booking, report upload, doctor notes, prescriptions, and medicine options.",
        "Pilot: secure backend, role-based access, audit logs, OCR extraction, hospital-specific configuration, clinician-reviewed prompts, and pilot analytics.",
        "Scale: hospital integrations, doctor availability APIs, ABDM-readiness, pharmacy integrations, analytics dashboards, clinician feedback loops, and offline knowledge-base fallback.",
    ]:
        add_bullet(doc, item)

    add_page(doc, "Clinical Safety, Compliance, and Governance")
    add_para(
        doc,
        "Healthcare AI requires a stronger governance story than ordinary SaaS. The safest commercialization path is to define the product as a support and routing layer under hospital governance, with doctors retaining responsibility for diagnosis, investigations, prescriptions, and treatment. The startup should avoid implying independent medical-device functionality until regulatory, clinical, and legal review supports that claim.",
        size=10.8,
    )
    add_matrix(
        doc,
        ["Risk", "Mitigation to build", "Investor/DPIIT implication"],
        [
            ("Misdiagnosis or overclaiming", "Use non-diagnostic language, emergency escalation, confidence limits, and doctor confirmation.", "Shows responsible innovation rather than reckless automation."),
            ("Sensitive health data", "Consent, encryption, access controls, retention policy, deletion/export, and DPDP-aligned processing.", "Improves hospital trust and enterprise readiness."),
            ("Model hallucination", "Structured outputs, red-flag rules outside the model, clinical prompt review, monitored fallback, and doctor feedback.", "Creates a credible AI governance moat."),
            ("Unsafe prescriptions", "Never generate prescriptions for patients; only store doctor-entered prescription content.", "Keeps clinical accountability with licensed professionals."),
            ("Dataset misuse", "Use approved cases only after de-identification, clinician review, versioning, and quality thresholds.", "Prepares for scalable local knowledge while protecting patients."),
        ],
        [2100, 4350, 2910],
    )
    add_heading(doc, "Production readiness items", 2)
    for item in [
        "Replace browser-local storage with secure backend persistence.",
        "Add session expiry, MFA for doctors/admins, authorization by role, and audit logs for every case access.",
        "Add patient consent screens for report upload, case learning, prescription storage, and pharmacy comparison.",
        "Add OCR and structured lab-value extraction with abnormal-range flagging and source verification.",
        "Document clinical review SOPs, model evaluation metrics, escalation policy, and incident response.",
    ]:
        add_bullet(doc, item)

    add_page(doc, "Market, Customers, and Go-To-Market")
    add_heading(doc, "Primary customer segments", 2)
    for item in [
        "Mid-size hospitals and specialty clinics seeking OPD intake automation and better patient routing.",
        "Hospital chains wanting configurable digital front-door workflows across multiple locations.",
        "Telemedicine, insurance, and care-navigation providers needing multilingual triage and doctor handoff.",
        "Corporate health programs and clinics managing employee care, reports, appointments, and follow-up.",
        "Government, NGO, and public-health pilots focused on accessible language support and early care navigation.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "GTM motion", 2)
    for item in [
        "Start with 2-3 hospital/clinic pilots in high-volume OPD categories such as fever, respiratory, gastro, UTI, diabetes, and general medicine.",
        "Sell as a digital front-door module: intake, routing, booking, and doctor handoff, not as a standalone chatbot.",
        "Use hospital-branded deployment widgets and configurable departments/doctors to reduce integration friction.",
        "Measure operational outcomes: reduced misrouting, shorter intake time, better report availability, more complete doctor notes, and appointment conversion.",
        "Convert pilots into annual SaaS contracts with setup fees, support fees, and integration expansion.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Expansion thesis", 2)
    for item in [
        "Department-specific assistants: dermatology, pediatrics, gynecology, pulmonology, diabetes, emergency pre-triage, and chronic care.",
        "ABDM-ready patient record and consent integrations where hospital strategy supports it.",
        "Doctor-side workflow intelligence: recommended reports, repeated case patterns, missed follow-up detection, and patient education materials.",
        "Marketplace/API potential for verified pharmacy pricing, diagnostics, second opinions, and care packages.",
    ]:
        add_bullet(doc, item)

    add_page(doc, "Business Model and Funding Case")
    add_heading(doc, "Revenue model", 2)
    for item in [
        "SaaS subscription per hospital, clinic, location, doctor group, or monthly patient-intake volume.",
        "Implementation and configuration fees for hospital-specific departments, doctors, slots, language, and workflows.",
        "Integration revenue for HIS/EMR, appointment, ABDM, OCR, pharmacy, diagnostics, and notification systems.",
        "Usage-based AI review charges where hospitals choose LLM-backed triage rather than rules-only flows.",
        "Analytics and quality dashboards for hospital operations, case completion, routing quality, and care-navigation metrics.",
        "Partner/referral revenue only where legally compliant, transparent, and clinically appropriate.",
    ]:
        add_bullet(doc, item)
    add_kv_table(
        doc,
        [
            ("Suggested round", "[Pre-seed / seed amount to be inserted after founder review]"),
            ("Use of funds", "Secure backend, clinician review, OCR/report parsing, model evaluation, hospital integrations, pilot deployments, sales, compliance, and customer success."),
            ("Milestones", "Complete product-grade MVP, launch pilots, validate safety metrics, convert paid customers, build hospital integration playbook, and create clinical governance evidence."),
            ("Investor thesis", "Large healthcare navigation need, AI-enabled workflow wedge, recurring SaaS model, doctor feedback data asset, and India-language healthcare access value."),
        ],
        widths=(2100, 7260),
    )

    add_page(doc, "Traction Evidence and Metrics")
    add_para(
        doc,
        "DPIIT and investors respond well to evidence, but healthcare evidence must be precise. If pilots are early, present product execution and signed intent honestly instead of overstating clinical validation.",
        size=10.8,
    )
    add_heading(doc, "Evidence to attach or reference", 2)
    for item in [
        "[ ] Product screenshots showing patient intake, triage, follow-up, booking, doctor queue, prescription, and medicine options.",
        "[ ] Docker/demo link or screen recording of the product flow.",
        "[ ] Architecture note showing modular configuration, backend plan, data model, LLM safety wrapper, and audit trail plan.",
        "[ ] Pilot LOIs, clinic/hospital interest emails, partner discussions, or signed proof-of-concept scope.",
        "[ ] Clinical advisor/doctor review notes for triage guardrails, red flags, and follow-up question quality.",
        "[ ] Founder profiles, technical roadmap, source-code repository evidence, and development timeline.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Metrics investors will ask for", 2)
    add_matrix(
        doc,
        ["Category", "Metrics to track"],
        [
            ("Product usage", "Cases created, triage runs, completed follow-ups, report uploads, appointments booked, doctor case views."),
            ("Operational value", "Intake time saved, routing correction rate, no-show reduction, consultation readiness, doctor note completeness."),
            ("Safety quality", "Red-flag recall, model fallback rate, low-confidence cases, doctor override rate, hallucination incidents, complaint rate."),
            ("Commercial", "Pilot conversion, ACV, setup cost, gross margin, support tickets, integration time, churn risk, expansion revenue."),
            ("Learning loop", "Clinician-approved cases, rejected cases, specialty coverage, report extraction accuracy, prompt/model version performance."),
        ],
        [2200, 7160],
    )

    add_page(doc, "Impact, Jobs, and Wealth Creation")
    add_heading(doc, "Employment generation potential", 2)
    for item in [
        "Product and engineering roles for SaaS, AI workflows, backend systems, security, integrations, speech/OCR, and analytics.",
        "Clinical operations roles for doctor review coordination, prompt validation, safety monitoring, and healthcare content governance.",
        "Implementation and customer success roles for hospital onboarding, department configuration, staff training, and pilot support.",
        "Sales/channel roles for hospitals, clinics, telemedicine partners, diagnostics networks, and enterprise health programs.",
        "Indirect jobs through medical coordinators, diagnostics partners, pharmacy integrations, health camps, and implementation consultants.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Public and economic value", 2)
    for item in [
        "Improves patient navigation by helping users reach the right specialist earlier.",
        "Supports Indian-language access through Hindi and Hinglish symptom intake.",
        "Reduces operational friction for hospitals by structuring symptoms, reports, and follow-up answers before consultation.",
        "Creates a reusable healthcare workflow platform that can be deployed across multiple hospitals and specialties.",
        "Builds data assets and IP around safe triage workflows, doctor feedback, routing logic, report summarization, and care-navigation analytics.",
    ]:
        add_bullet(doc, item)

    add_page(doc, "DPIIT Submission Checklist")
    add_heading(doc, "Must-have application materials", 2)
    for item in [
        "[ ] Certificate of incorporation / registration certificate.",
        "[ ] PAN of the entity and authorized signatory details.",
        "[ ] Brief write-up on innovation, scalability, employment generation, and wealth creation potential.",
        "[ ] Pitch deck or support PDF covering problem, solution, market, business model, traction, team, and impact.",
        "[ ] Website/product link, app screenshots, demo video, prototype evidence, or technical documentation.",
        "[ ] Board resolution/authorization letter if someone applies on behalf of the company.",
        "[ ] IP/patent/trademark details if applicable; otherwise note proprietary software, workflows, data structures, prompts, and roadmap.",
        "[ ] Customer/pilot proof, LOIs, clinical advisor notes, invoices, testimonials, or hospital discussions where available.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Quality checks before upload", 2)
    for item in [
        "Use only verified facts; do not claim hospitals, doctors, approvals, revenue, pilots, certifications, or clinical validation unless evidence is attached.",
        "Keep company name, brand name, founder names, incorporation date, and registration numbers consistent across all documents.",
        "Avoid presenting the product as a diagnostic or prescription system; position it as triage, routing, workflow, and doctor-support infrastructure.",
        "Confirm current DPIIT criteria and required upload fields on the official Startup India portal immediately before submission.",
        "Use a clear file name such as Axnovus_Care_Agent_DPIIT_Startup_India_Investor_Document.pdf.",
    ]:
        add_bullet(doc, item)

    add_page(doc, "One-Page Investor Narrative")
    add_para(
        doc,
        "Axnovus Care Agent is building the AI workflow layer for Indian healthcare access and OPD coordination. The problem starts before diagnosis: patients describe symptoms in mixed language, may not know which doctor to book, often carry reports that are not interpreted early, and doctors receive incomplete handoffs. Hospitals need a digital front door that listens, structures, asks, routes, and hands off responsibly.",
        size=10.8,
    )
    add_para(
        doc,
        "The product begins as a configurable patient-to-doctor workflow: symptom intake by text or Hindi speech, report upload, possible-condition triage, follow-up questions, specialist routing, appointment booking, doctor review, prescription capture, and medicine option comparison. It uses LLM reasoning where appropriate, but keeps red flags, clinical safety, doctor confirmation, and auditability central to the design.",
        size=10.8,
    )
    add_para(
        doc,
        "The business can scale through hospital SaaS subscriptions, setup and integration fees, usage-based AI review, workflow analytics, and healthcare partner integrations. The data advantage compounds through clinician-reviewed cases, doctor corrections, report requests, and routing outcomes, eventually allowing a safer local knowledge layer and stronger specialty-specific workflows.",
        size=10.8,
    )
    add_heading(doc, "Immediate funding priorities", 2)
    for item in [
        "Move from prototype to secure product-grade MVP with backend persistence, role-based access, audit logs, and consent.",
        "Run 2-3 controlled hospital/clinic pilots with doctor review and measurable operational outcomes.",
        "Build OCR/report extraction, clinical evaluation datasets, and model safety monitoring.",
        "Validate commercial packaging: setup fee, subscription price, AI usage cost, support burden, and integration timeline.",
        "Prepare DPIIT, accelerator, grant, and seed investor materials with verified evidence.",
    ]:
        add_bullet(doc, item)

    add_page(doc, "Founder Review Notes")
    add_para(
        doc,
        "This document is intentionally structured for both DPIIT recognition and investor review. Before external use, replace all bracketed fields, attach factual evidence, and have legal, clinical, and data-privacy assumptions reviewed. The strongest submission will pair this narrative with product screenshots, incorporation documents, pilot proof, a short demo link, and a technical architecture note.",
        size=10.8,
    )
    add_heading(doc, "Open items to finalize", 2)
    for item in [
        "[ ] Legal entity name, registration number, incorporation date, and address.",
        "[ ] Founder names, bios, roles, shareholding, and relevant AI/healthcare/business experience.",
        "[ ] Current product status: prototype, MVP, pilot, paid pilot, or revenue stage.",
        "[ ] Existing customers, hospital discussions, doctors/advisors, pilots, LOIs, grants, or deployments.",
        "[ ] Funding amount, proposed use of funds, 12-18 month budget, and hiring plan.",
        "[ ] Screenshots, demo URL, product architecture visual, clinical safety note, and IP/trademark status.",
        "[ ] Confirmation of current DPIIT eligibility thresholds and application fields before upload.",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Founder review stance",
        "Keep the story ambitious but disciplined. DPIIT and investor reviewers will trust the document more if it clearly separates current prototype capability, pilot-stage evidence, product roadmap, and regulated healthcare claims.",
        fill=GRAY_FILL,
    )

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
